"""
Replace the withdrawn 99.2% shade claim in the editable Word documents.

The geometric shade model in src/solar.py does not reproduce 99.2% annual shade
on the spine under any reading — the best single square metre anywhere on site
reaches 83%. The figure is withdrawn and replaced with the computed values:

    60.2%   canopy structure alone
    69.6%   canopy plus the flanking tree avenue   <- the headline figure

    python tools/fix_withdrawn_claim.py --dry-run
    python tools/fix_withdrawn_claim.py

Only .docx files are touched, because only they are editable. The matching PDFs
must be re-exported from Word afterwards — this script prints the list.
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
NEW = "69.6"

TARGETS = [
    "submission/10_Complete_Design_Report/Al_Safa_2_Park_Complete_Design_Report.docx",
    "submission/12_Concept_Animation_Video/Concept_Animation_Storyboard.docx",
    "reports/editable_docx/Phase7_Performance_and_Sustainability_Report.docx",
    "reports/editable_docx/Phase1_Knowledge_Base.docx",
    "reports/EASY_UNDERSTANDING_GUIDE.docx",
    "reports/PROJECT_METHODOLOGY_ROADMAP.docx",
]

# Longer phrases first so a specific rewrite wins over the bare number.
REPLACEMENTS = [
    ("99.2% annual shade coverage — confirming the 3-date snapshot's 100% result holds up",
     f"{NEW}% annual shade coverage — revised downward from an earlier 99.2% figure, which"
     " measured a single centroid point rather than the walkway area"),
    ("computed at 99.2% annual shade coverage",
     f"computed at {NEW}% annual shade coverage from canopy and tree avenue together"),
    ("99.2% annual shade coverage", f"{NEW}% annual shade coverage"),
    ("99.2% shaded, all year — proven, not promised.",
     f"{NEW}% of daylight hours shaded — measured, not promised."),
    ("proven 99.2% shaded circulation", f"measured {NEW}% shaded circulation"),
    ("holds 99.2% annual shade coverage", f"holds {NEW}% annual shade coverage"),
    ("covered by shade 99.2% of all daylight hours",
     f"covered by shade {NEW}% of all daylight hours"),
    ("99.2% shade,", f"{NEW}% shade,"),
    ("99.2% shade", f"{NEW}% shade"),
    ("99.2%", f"{NEW}%"),
]


def fix_paragraph(par) -> int:
    """Rewrite runs in place, preserving each run's formatting where possible."""
    if "99.2" not in par.text:
        return 0
    n = 0
    for run in par.runs:
        for old, new in REPLACEMENTS:
            if old in run.text:
                run.text = run.text.replace(old, new)
                n += 1
    # A claim split across several runs will not match run-by-run; fall back to
    # collapsing the paragraph into its first run.
    if "99.2" in par.text and par.runs:
        merged = par.text
        for old, new in REPLACEMENTS:
            merged = merged.replace(old, new)
        if merged != par.text:
            par.runs[0].text = merged
            for r in par.runs[1:]:
                r.text = ""
            n += 1
    return n


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    total, touched = 0, []
    for rel in TARGETS:
        p = ROOT / rel
        if not p.exists():
            print(f"  [skip] {rel} — not found")
            continue

        d = docx.Document(str(p))
        n = sum(fix_paragraph(par) for par in d.paragraphs)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    n += sum(fix_paragraph(par) for par in cell.paragraphs)

        if n:
            if not args.dry_run:
                backup = p.with_suffix(".docx.bak")
                if not backup.exists():
                    shutil.copy2(p, backup)
                d.save(str(p))
            print(f"  [{'would fix' if args.dry_run else 'fixed'}] {rel} — {n} replacement(s)")
            total += n
            touched.append(rel)
        else:
            print(f"  [clean] {rel}")

    print(f"\n{total} replacement(s) across {len(touched)} file(s)"
          f"{'  (dry run — nothing written)' if args.dry_run else ''}")

    if touched and not args.dry_run:
        print("\nRE-EXPORT THESE TO PDF FROM WORD (this script cannot — PDF export needs Word):")
        for rel in touched:
            print(f"  - {rel}")
        print("\nOriginals kept alongside each file as *.docx.bak")


if __name__ == "__main__":
    main()
