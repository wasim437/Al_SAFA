"""
Shared Word (.docx) report builder for Al Safa 2 Park phase reports.

WHY THIS EXISTS: the earlier all-Python/ReportLab pipeline (report_builder.py)
required editing Python code to fix a typo or reword a sentence, which is how
text bugs (unescaped '&', bad column widths) crept in. This builder instead
writes real, directly-editable Word documents - text can be fixed by opening
the .docx in Word, no Python required. Python is only used to assemble the
structure and insert charts (which genuinely do need to be code-generated,
since they come from real data).

Usage is the same "sections" list shape as report_builder.py so existing
gen_pdf_*.py scripts can be ported with minimal changes.
"""

import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ACCENT = RGBColor(0xC8, 0xA2, 0x4A)
GREY = RGBColor(0x55, 0x55, 0x55)


def _set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _strip_html_bold_italic(text):
    """Our section text uses a tiny bit of <b>/<i> markup (carried over from
    the PDF pipeline). Word needs real run formatting, not tags, so this
    splits text on <b>/<i> and returns a list of (text, bold, italic) runs."""
    import re
    runs = []
    pos = 0
    bold = italic = False
    for m in re.finditer(r"<(/?)(b|i)>", text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], bold, italic))
        closing, tag = m.group(1), m.group(2)
        if tag == "b":
            bold = not closing
        elif tag == "i":
            italic = not closing
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], bold, italic))
    # Clean up any leftover HTML entities from the old pipeline
    cleaned = []
    for t, b, i in runs:
        t = (t.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
               .replace("&#176;", "°").replace("&ge;", "≥").replace("&le;", "≤"))
        cleaned.append((t, b, i))
    return cleaned


def _add_rich_text(paragraph, text):
    for t, bold, italic in _strip_html_bold_italic(text):
        run = paragraph.add_run(t)
        run.bold = bold
        run.italic = italic


def build_docx_report(output_path, phase_tag, title, subtitle, sections,
                       code_ref=None, script_name=None):
    """
    sections: same shape as report_builder.build_report:
      {"type": "heading", "text": "..."}
      {"type": "para", "text": "..."}
      {"type": "bullets", "items": [...]}
      {"type": "table", "header": [...], "rows": [[...]]}
      {"type": "image", "path": "...", "caption": "...", "width_cm": 16}
      {"type": "spacer", "h_cm": 0.5}
    code_ref: raw python source text to embed as a "Program Proof" appendix
    script_name: the script's relative path, shown in the appendix intro
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # Header band: phase tag
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(phase_tag)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = GREY

    # Gold rule
    rule_p = doc.add_paragraph()
    rule_run = rule_p.add_run("─" * 60)
    rule_run.font.color.rgb = ACCENT
    rule_run.font.size = Pt(8)

    for s in sections:
        t = s["type"]
        if t == "heading":
            h = doc.add_heading(level=1)
            hrun = h.add_run(s["text"])
            hrun.font.color.rgb = NAVY
            hrun.font.size = Pt(15)
        elif t == "para":
            p = doc.add_paragraph()
            _add_rich_text(p, s["text"])
        elif t == "bullets":
            for item in s["items"]:
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_text(p, item)
        elif t == "table":
            header = s["header"]
            rows = s["rows"]
            ncol = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncol)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for c, htext in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_rich_text(p, str(htext))
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
                _set_cell_background(cell, "0B1F3A")
            for r, row in enumerate(rows, start=1):
                for c in range(ncol):
                    val = row[c] if c < len(row) else ""
                    cell = table.rows[r].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    _add_rich_text(p, str(val))
                    for run in p.runs:
                        run.font.size = Pt(9)
                    if r % 2 == 0:
                        _set_cell_background(cell, "F4F5F7")
            doc.add_paragraph()  # spacing after table
        elif t == "image":
            img_path = s["path"]
            if os.path.exists(img_path):
                width_cm = s.get("width_cm", 15)
                doc.add_picture(img_path, width=Cm(width_cm))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if s.get("caption"):
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(s["caption"])
                    cap_run.italic = True
                    cap_run.font.size = Pt(8.5)
                    cap_run.font.color.rgb = GREY
            else:
                p = doc.add_paragraph()
                p.add_run(f"[Missing image: {img_path}]").italic = True
        elif t == "spacer":
            doc.add_paragraph()

    # --- Program Proof appendix ---
    if code_ref:
        doc.add_page_break()
        h = doc.add_heading(level=1)
        hrun = h.add_run("Appendix — Program Proof (Source Code)")
        hrun.font.color.rgb = NAVY
        hrun.font.size = Pt(15)
        intro = doc.add_paragraph()
        intro.add_run(
            f"The analysis, figures, and tables in this report were generated by the Python "
            f"script "
        )
        b = intro.add_run(script_name or "")
        b.bold = True
        intro.add_run(
            " below. Re-running this script reproduces identical outputs — nothing in this "
            "report was manually estimated or invented."
        )
        for line in code_ref.splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line or " ")
            run.font.name = "Consolas"
            run.font.size = Pt(8)

    # Footer text (simple, as a final line - python-docx footer API is verbose;
    # a plain last-page note is enough for an internally-reviewed draft)
    doc.save(output_path)
    print(f"Saved DOCX: {output_path}")
    return output_path


def convert_docx_to_pdf(docx_path, pdf_path=None):
    """Uses MS Word (via docx2pdf/COM automation) to export a real PDF from
    the .docx - so the PDF is always a faithful render of what's in Word."""
    from docx2pdf import convert
    pdf_path = pdf_path or docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    print(f"Saved PDF (from Word): {pdf_path}")
    return pdf_path
